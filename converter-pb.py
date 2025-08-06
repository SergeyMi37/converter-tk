# Вот модифицированный код с добавлением прогрессбара для отображения хода конвертации:

# ```python
import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from docx import Document
import pdfkit
import pypandoc
from htmldocx import HtmlToDocx
from striprtf.striprtf import rtf_to_text
import markdown

class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Конвертер документов")
        
        # Переменные для хранения путей
        self.source_dir = tk.StringVar()
        self.target_dir = tk.StringVar()
        self.conversion_mode = tk.StringVar(value="rtf -> docx")
        
        # Привязываем событие закрытия по клавише Esc
        self.root.bind('<Escape>', lambda event: self.root.destroy())        
        
        # Создаем интерфейс
        self.create_widgets()
        
    def create_widgets(self):
        # Поля для выбора директорий
        tk.Label(self.root, text="Исходная директория:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self.root, textvariable=self.source_dir, width=50).grid(row=0, column=1, padx=5, pady=5)
        tk.Button(self.root, text="Обзор...", command=self.browse_source).grid(row=0, column=2, padx=5, pady=5)
        
        tk.Label(self.root, text="Целевая директория:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self.root, textvariable=self.target_dir, width=50).grid(row=1, column=1, padx=5, pady=5)
        tk.Button(self.root, text="Обзор...", command=self.browse_target).grid(row=1, column=2, padx=5, pady=5)
        
        # Выбор режима конвертации
        tk.Label(self.root, text="Режим конвертации:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        modes = [
            "rtf -> docx",
            "docx -> pdf",
            "rtf -> pdf",
            "html -> docx",
            "html -> rtf",
            "html -> pdf"
        ]
        tk.OptionMenu(self.root, self.conversion_mode, *modes).grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        # Прогрессбар
        self.progress = ttk.Progressbar(self.root, orient="horizontal", length=400, mode="determinate")
        self.progress.grid(row=3, column=0, columnspan=3, padx=5, pady=5)
        
        # Метка для отображения текущего файла
        self.current_file_label = tk.Label(self.root, text="", fg="blue")
        self.current_file_label.grid(row=4, column=0, columnspan=3, padx=5, pady=5)
        
        # Кнопки
        button_frame = tk.Frame(self.root)
        button_frame.grid(row=5, column=0, columnspan=3, pady=10)
        
        tk.Button(button_frame, text="Конвертировать", command=self.convert_files, bg="green", fg="white").pack(side="left", padx=10)
        tk.Button(button_frame, text="Закрыть", command=self.root.destroy, bg="red", fg="white").pack(side="right", padx=10)

    def browse_source(self):
        directory = filedialog.askdirectory()
        if directory:
            self.source_dir.set(directory)
    
    def browse_target(self):
        directory = filedialog.askdirectory()
        if directory:
            self.target_dir.set(directory)

    def update_progress(self, value, filename=""):
        self.progress['value'] = value
        self.current_file_label.config(text=f"Обработка: {filename}" if filename else "")
        self.root.update_idletasks()

    def convert_files(self):
        source = self.source_dir.get()
        target = self.target_dir.get()
        mode = self.conversion_mode.get()
        
        if not source or not target:
            messagebox.showerror("Ошибка", "Выберите исходную и целевую директории")
            return
        
        try:
            # Создаем целевую директорию, если ее нет
            os.makedirs(target, exist_ok=True)
            
            # Получаем список файлов для обработки
            files = []
            if mode.startswith("rtf"):
                files = [f for f in os.listdir(source) if f.lower().endswith('.rtf')]
            elif mode.startswith("docx"):
                files = [f for f in os.listdir(source) if f.lower().endswith('.docx')]
            elif mode.startswith("html"):
                files = [f for f in os.listdir(source) if f.lower().endswith('.html')]
            
            if not files:
                messagebox.showwarning("Предупреждение", "Нет файлов для конвертации в выбранной директории")
                return
            
            # Настраиваем прогрессбар
            self.progress['maximum'] = len(files)
            self.progress['value'] = 0
            self.current_file_label.config(text="")
            
            # Обрабатываем файлы
            for i, filename in enumerate(files, 1):
                self.update_progress(i, filename)
                
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0])
                
                if mode == "rtf -> docx":
                    output_file += '.docx'
                    self.convert_rtf_to_docx(input_file, output_file)
                elif mode == "docx -> pdf":
                    output_file += '.pdf'
                    self.convert_docx_to_pdf(input_file, output_file)
                elif mode == "rtf -> pdf":
                    output_file += '.pdf'
                    self.convert_rtf_to_pdf(input_file, output_file)
                elif mode == "html -> docx":
                    output_file += '.docx'
                    self.convert_html_to_docx(input_file, output_file)
                elif mode == "html -> rtf":
                    output_file += '.rtf'
                    self.convert_html_to_rtf(input_file, output_file)
                elif mode == "html -> pdf":
                    output_file += '.pdf'
                    self.convert_html_to_pdf(input_file, output_file)
            
            messagebox.showinfo("Успех", f"Конвертация завершена успешно!\nОбработано файлов: {len(files)}")
            self.update_progress(0)  # Сбрасываем прогрессбар
            
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")
            self.update_progress(0)  # Сбрасываем прогрессбар
    
    def convert_rtf_to_docx(self, input_file, output_file):
        """Конвертация RTF в DOCX с использованием python-docx"""
        with open(input_file, 'r', encoding='utf-8') as f:
            rtf_text = f.read()
        
        plain_text = rtf_to_text(rtf_text)
        doc = Document()
        doc.add_paragraph(plain_text)
        doc.save(output_file)
    
    def convert_docx_to_pdf(self, input_file, output_file):
        """Конвертация DOCX в PDF с использованием pdfkit"""
        doc = Document(input_file)
        html_content = "\n".join([p.text for p in doc.paragraphs])
        
        options = {
            'encoding': 'UTF-8',
            'quiet': ''
        }
        
        pdfkit.from_string(html_content, output_file, options=options)
    
    def convert_rtf_to_pdf(self, input_file, output_file):
        """Конвертация RTF в PDF через промежуточное преобразование в текст"""
        with open(input_file, 'r', encoding='utf-8') as f:
            rtf_text = f.read()
        
        plain_text = rtf_to_text(rtf_text)
        
        options = {
            'encoding': 'UTF-8',
            'quiet': ''
        }
        pdfkit.from_string(plain_text, output_file, options=options)
    
    def convert_html_to_docx(self, input_file, output_file):
        """Конвертация HTML в DOCX с использованием htmldocx"""
        parser = HtmlToDocx()
        
        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        doc = parser.parse_html_string(html_content)
        doc.save(output_file)
    
    def convert_html_to_rtf(self, input_file, output_file):
        """Конвертация HTML в RTF через промежуточное преобразование в текст"""
        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        plain_text = html_content.replace('<', ' <').replace('>', '> ')
        plain_text = ' '.join(plain_text.split())
        
        rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}" + "\n"
        rtf_content += r"{\colortbl;\red0\green0\blue0;}" + "\n"
        rtf_content += plain_text + "\n}"
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(rtf_content)
    
    def convert_html_to_pdf(self, input_file, output_file):
        """Конвертация HTML в PDF с использованием pdfkit"""
        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        options = {
            'encoding': 'UTF-8',
            'quiet': ''
        }
        pdfkit.from_string(html_content, output_file, options=options)

if __name__ == "__main__":
    # Проверяем зависимости
    try:
        import pdfkit
        import pypandoc
        from htmldocx import HtmlToDocx
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        print("Установите необходимые зависимости:")
        print("pip install pdfkit pypandoc python-docx htmldocx striprtf")
        exit(1)
    
    # Проверяем установлен ли wkhtmltopdf для pdfkit
    try:
        pdfkit.from_string('test', 'doc/test.pdf')
    except OSError:
        print("Установите wkhtmltopdf для работы с PDF:")
        print("Скачайте с https://wkhtmltopdf.org/downloads.html")
        print("И добавьте в PATH или укажите путь в config:")
        print("pdfkit.configuration(wkhtmltopdf='/path/to/wkhtmltopdf')")
        exit(1)
    
    root = tk.Tk()
    app = FileConverterApp(root)
    root.mainloop()
# ```

# Основные изменения:

# 1. Добавлен прогрессбар (`ttk.Progressbar`) для визуализации хода конвертации
# 2. Добавлена метка (`current_file_label`) для отображения имени текущего обрабатываемого файла
# 3. Добавлен метод `update_progress()` для обновления прогрессбара и метки
# 4. Модифицирован метод `convert_files()` для работы с прогрессбаром:
#    - Подсчет общего количества файлов для обработки
#    - Пошаговое обновление прогрессбара
#    - Отображение текущего файла
# 5. Улучшено сообщение о завершении конвертации с указанием количества обработанных файлов
# 6. Добавлен сброс прогрессбара при ошибке или завершении

# Теперь при конвертации пользователь видит:
# - Полосу прогресса, заполняющуюся по мере обработки файлов
# - Имя текущего обрабатываемого файла
# - Общее количество файлов (по длине прогрессбара)