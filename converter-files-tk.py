import os
import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import pdfkit
import pypandoc
from htmldocx import HtmlToDocx
from striprtf.striprtf import rtf_to_text
import markdown

class FileConverterApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Конвертер документов")
        
        # Переменные для хранения путей
        self.source_dir = tk.StringVar()
        self.target_dir = tk.StringVar()
        self.conversion_mode = tk.StringVar(value="rtf -> docx")
        
        # Создаем интерфейс
        self.create_widgets()
        
    def create_widgets(self):
        # Поля для выбора директорий
        tk.Label(self.root, text="Исходная директория:").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self.root, textvariable=self.source_dir, width=50).grid(row=0, column=1, padx=5, pady=5)
        tk.Button(self.root, text="Обзор...", command=self.browse_source).grid(row=0, column=2, padx=5, pady=5)
        
        tk.Label(self.root, text="Целевая директория:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self.root, textvariable=self.target_dir, width=50).grid(row=1, column=1, padx=5, pady=5)
        tk.Button(self.root, text="Обзор...", command=self.browse_target).grid(row=1, column=2, padx=5, pady=5)
        
        # Выбор режима конвертации
        tk.Label(self.root, text="Режим конвертации:").grid(row=2, column=0, padx=5, pady=5, sticky="w")
        modes = [
            "rtf -> docx",
            "docx -> pdf",
            "rtf -> pdf",
            "html -> docx",
            "html -> rtf"
        ]
        tk.OptionMenu(self.root, self.conversion_mode, *modes).grid(row=2, column=1, padx=5, pady=5, sticky="w")
        
        # Кнопка конвертации
        tk.Button(self.root, text="Конвертировать", command=self.convert_files, bg="green", fg="white").grid(row=3, column=1, pady=10)
        
    def browse_source(self):
        directory = filedialog.askdirectory()
        if directory:
            self.source_dir.set(directory)
    
    def browse_target(self):
        directory = filedialog.askdirectory()
        if directory:
            self.target_dir.set(directory)
    
    def convert_files(self):
        source = self.source_dir.get()
        target = self.target_dir.get()
        mode = self.conversion_mode.get()
        
        if not source or not target:
            messagebox.showerror("Ошибка", "Выберите исходную и целевую директории")
            return
        
        try:
            # Создаем целевую директорию, если ее нет
            os.makedirs(target, exist_ok=True)
            
            # Обрабатываем файлы в зависимости от режима
            if mode == "rtf -> docx":
                self.convert_rtf_to_docx(source, target)
            elif mode == "docx -> pdf":
                self.convert_docx_to_pdf(source, target)
            elif mode == "rtf -> pdf":
                self.convert_rtf_to_pdf(source, target)
            elif mode == "html -> docx":
                self.convert_html_to_docx(source, target)
            elif mode == "html -> rtf":
                self.convert_html_to_rtf(source, target)
                
            messagebox.showinfo("Успех", "Конвертация завершена успешно!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Произошла ошибка: {str(e)}")
    
    def convert_rtf_to_docx(self, source, target):
        """Конвертация RTF в DOCX с использованием python-docx"""
        for filename in os.listdir(source):
            if filename.lower().endswith('.rtf'):
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0] + '.docx')
                
                # Читаем RTF файл
                with open(input_file, 'r', encoding='utf-8') as f:
                    rtf_text = f.read()
                
                # Конвертируем RTF в простой текст
                plain_text = rtf_to_text(rtf_text)
                
                # Создаем DOCX документ
                doc = Document()
                doc.add_paragraph(plain_text)
                doc.save(output_file)
    
    def convert_docx_to_pdf(self, source, target):
        """Конвертация DOCX в PDF с использованием pdfkit"""
        for filename in os.listdir(source):
            if filename.lower().endswith('.docx'):
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0] + '.pdf')
                
                # Конвертируем docx в html, затем в pdf
                doc = Document(input_file)
                html_content = "\n".join([p.text for p in doc.paragraphs])
                
                options = {
                    'encoding': 'UTF-8',
                    'quiet': ''
                }
                
                pdfkit.from_string(html_content, output_file, options=options)
    
    def convert_rtf_to_pdf(self, source, target):
        """Конвертация RTF в PDF через промежуточное преобразование в текст"""
        for filename in os.listdir(source):
            if filename.lower().endswith('.rtf'):
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0] + '.pdf')
                
                # Читаем RTF файл
                with open(input_file, 'r', encoding='utf-8') as f:
                    rtf_text = f.read()
                
                # Конвертируем RTF в простой текст
                plain_text = rtf_to_text(rtf_text)
                
                # Сохраняем в PDF
                options = {
                    'encoding': 'UTF-8',
                    'quiet': ''
                }
                pdfkit.from_string(plain_text, output_file, options=options)
    
    def convert_html_to_docx(self, source, target):
        """Конвертация HTML в DOCX с использованием htmldocx"""
        parser = HtmlToDocx()
        
        for filename in os.listdir(source):
            if filename.lower().endswith('.html'):
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0] + '.docx')
                
                # Читаем HTML файл
                with open(input_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Конвертируем в DOCX
                doc = parser.parse_html_string(html_content)
                doc.save(output_file)
    
    def convert_html_to_rtf(self, source, target):
        """Конвертация HTML в RTF через промежуточное преобразование в текст"""
        for filename in os.listdir(source):
            if filename.lower().endswith('.html'):
                input_file = os.path.join(source, filename)
                output_file = os.path.join(target, os.path.splitext(filename)[0] + '.rtf')
                
                # Читаем HTML файл
                with open(input_file, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                
                # Конвертируем HTML в простой текст (упрощенный вариант)
                plain_text = html_content.replace('<', ' <').replace('>', '> ')
                plain_text = ' '.join(plain_text.split())
                
                # Сохраняем как RTF (упрощенный RTF)
                rtf_content = r"{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}" + "\n"
                rtf_content += r"{\colortbl;\red0\green0\blue0;}" + "\n"
                rtf_content += plain_text + "\n}"
                
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(rtf_content)

if __name__ == "__main__":
    # Проверяем зависимости
    try:
        import pdfkit
        import pypandoc
        from htmldocx import HtmlToDocx
        from striprtf.striprtf import rtf_to_text
    except ImportError:
        print("Установите необходимые зависимости:")
        print("pip install pdfkit pypandoc python-docx htmldocx striprtf")
        exit(1)
    
    # Проверяем установлен ли wkhtmltopdf для pdfkit
    try:
        pdfkit.from_string('test', 'test.pdf')
    except OSError:
        print("Установите wkhtmltopdf для работы с PDF:")
        print("Скачайте с https://wkhtmltopdf.org/downloads.html")
        print("И добавьте в PATH или укажите путь в config:")
        print("pdfkit.configuration(wkhtmltopdf='/path/to/wkhtmltopdf')")
        exit(1)
    
    root = tk.Tk()
    app = FileConverterApp(root)
    root.mainloop()